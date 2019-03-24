from docx import Document
import os


def build_table(doc_table):
    output = '<table border="1" cellspacing="0">\n<tbody>\n'
    for row in doc_table.rows:
        output += '<tr>'
        if row.cells[0].text == row.cells[1].text:
            output += '<td align="center" colspan="6">' + row.cells[0].text + '</td>'
        else:
            for i, cell in enumerate(row.cells):
                output += '<td align="center">' + cell.text + '</td>'
        output += '</tr>\n'
    output += '</tbody>\n</table>'
    return output


def comment_generator(doc_txt):
    comment = ''
    for line in doc_txt[1:]:
        comment += '<br/><p>' + line + '</p>'
    return comment


def build_html(doc_table, doc_txt):
    output = '<!DOCTYPE html>\n<html>\n<head>\n<meta charset="utf-8"/>\n</head>\n<body>\n'
    output += '<br/><p align="center">{}</p><br/>'.format(doc_txt[0])
    output += build_table(doc_table)
    output += comment_generator(doc_txt) + '</body>\n</html>'
    return output


print('welcome to brutal docx2html converter script\n'
      'this script search .docx files in /source folder \n'
      'and convert them to .html files without css-styles and other laces\n')
source_files = []
for file in os.listdir('source'):
    if file.endswith('.docx'):
        source_files.append(file)
        print('{0:>1}: {1:>40}'.format(len(source_files), file))

docxpath = ''
while not ((docxpath in source_files) or (docxpath.isdigit() and 0 < int(docxpath) <= len(source_files))):
    docxpath = input("\nenter filename 'filename.docx' for source file or its number from list above\n")
if docxpath.isdigit():
    docxpath = source_files[int(docxpath) - 1]

document = Document(os.path.join('source', docxpath))
table = document.tables[0]
txt = [s.text for s in document.paragraphs if s.text != '']

html = open(os.path.join('output', docxpath[:-4] + 'html'), 'w')
html.write(build_html(table, txt))
html.close()
print('\n.docx file converted successfully and saved to /output dir')
